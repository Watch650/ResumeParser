import fitz  # PyMuPDF
import docx  # python-docx
import os
import pytesseract
import pdf2image
import logging
import re
from typing import Optional, Dict, List, Union
from docx.document import Document

# Configure logging
logging.basicConfig(filename='cv_extraction.log', level=logging.INFO,
                   format='%(asctime)s - %(levelname)s - %(message)s')

# Section detection patterns
SECTION_PATTERNS = {
    'personal': re.compile(r'personal\s*details|information|profile', re.I),
    'summary': re.compile(r'professional\s*summary|objective|about', re.I),
    'skills': re.compile(r'skills|competencies|technologies', re.I),
    'experience': re.compile(r'experience|employment|work\s*history', re.I),
    'education': re.compile(r'education|qualifications|degrees', re.I),
    'projects': re.compile(r'projects|portfolio', re.I)
}

def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF using PyMuPDF."""
    try:
        doc = fitz.open(path)
        full_text = []
        for page in doc:
            text = page.get_text()
            if text:  # Only add if text was found
                full_text.append(text)
        return "\n".join(full_text).strip()
    except Exception as e:
        logging.error(f"Error reading PDF {path}: {str(e)}")
        return ""

def ocr_pdf(path: str, lang: str = "eng+vie+jpn", page_limit: Optional[int] = None) -> str:
    """Fallback OCR if text extraction fails."""
    try:
        images = pdf2image.convert_from_path(path)
        text = []
        for i, img in enumerate(images):
            if page_limit and i >= page_limit:
                break
            text.append(pytesseract.image_to_string(img, lang=lang))
        return "\n".join(text).strip()
    except Exception as e:
        logging.error(f"OCR failed for {path}: {str(e)}")
        return ""

def extract_tables(doc: Document) -> Dict[str, List[List[str]]]:
    """Extract and structure tables from DOCX."""
    tables_data = {}
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_data[f"table_{i+1}"] = table_data
    return tables_data

def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive whitespace and line breaks
    text = re.sub(r'\s+', ' ', text).strip()
    # Fix common OCR/table artifacts
    text = re.sub(r'(?<=\w) ,', ',', text)
    text = re.sub(r'\.\s+\.', '.', text)
    return text

def detect_section(text: str) -> Optional[str]:
    """Detect if text is a section header and return normalized section name."""
    text_lower = text.lower().strip(':* ')
    for section_name, pattern in SECTION_PATTERNS.items():
        if pattern.search(text_lower):
            return section_name
    return None

def extract_structured_cv(doc: Document) -> Dict[str, Union[List[str], Dict[str, List[List[str]]]]]:
    """Extract structured content from DOCX with section detection."""
    structured_data = {
        'metadata': {},
        'sections': {},
        'tables': {}
    }

    current_section = None
    current_content = []
    tables_data = extract_tables(doc)
    structured_data['tables'] = tables_data

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        section_match = detect_section(text)
        if section_match:
            if current_section and current_content:
                structured_data['sections'][current_section] = current_content
            current_section = section_match
            current_content = []
        else:
            if current_section:
                current_content.append(clean_text(text))
            else:
                if 'unstructured' not in structured_data['sections']:
                    structured_data['sections']['unstructured'] = []
                structured_data['sections']['unstructured'].append(clean_text(text))

    if current_section and current_content:
        structured_data['sections'][current_section] = current_content

    # Extract metadata from personal details if available
    if 'personal' in structured_data['sections']:
        for line in structured_data['sections']['personal']:
            if ':' in line:
                key, val = line.split(':', 1)
                structured_data['metadata'][key.strip().lower()] = val.strip()

    return structured_data

def format_structured_output(structured_data: Dict) -> str:
    """Convert structured data to formatted text output."""
    output = []

    # Add metadata
    if structured_data['metadata']:
        output.append("=== PERSONAL DETAILS ===")
        for key, val in structured_data['metadata'].items():
            output.append(f"{key.title()}: {val}")

    # Add sections
    for section, content in structured_data['sections'].items():
        if section == 'personal':  # Already handled in metadata
            continue
        output.append(f"\n=== {section.upper()} ===")
        output.extend(content)

    # Add tables
    if structured_data['tables']:
        output.append("\n=== TABLES ===")
        for table_name, table_data in structured_data['tables'].items():
            output.append(f"\nTable: {table_name}")
            for row in table_data:
                output.append(" | ".join(row))

    return "\n".join(output)

def extract_text_from_docx(path: str, structured: bool = True) -> Union[str, Dict]:
    """Extract text from DOCX with optional structured output."""
    try:
        doc = docx.Document(path)
        if structured:
            structured_data = extract_structured_cv(doc)
            return format_structured_output(structured_data)
        else:
            # Fallback to simple extraction
            return "\n".join(para.text for para in doc.paragraphs if para.text).strip()
    except Exception as e:
        logging.error(f"Error reading DOCX {path}: {str(e)}")
        return ""

def extract_text(file_path: str, ocr_lang: str = "eng+vie+jpn",
                attempt_ocr: bool = True, ocr_page_limit: Optional[int] = None,
                structured: bool = True) -> Union[str, Dict]:
    """
    Enhanced text extraction with structured output option.

    Args:
        file_path: Path to input file
        ocr_lang: Languages for OCR
        attempt_ocr: Whether to attempt OCR fallback
        ocr_page_limit: Limit for OCR pages
        structured: Return structured data if True

    Returns:
        Extracted text or structured data
    """
    if not os.path.exists(file_path):
        logging.error(f"File not found: {file_path}")
        return "" if structured else {}

    ext = os.path.splitext(file_path)[1].lower()
    result = "" if structured else {}

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
        if attempt_ocr and not text.strip():
            logging.info(f"Attempting OCR for {file_path}")
            text = ocr_pdf(file_path, lang=ocr_lang, page_limit=ocr_page_limit)
        result = text
    elif ext == ".docx":
        result = extract_text_from_docx(file_path, structured)
    else:
        logging.warning(f"Unsupported file type: {file_path}")

    return result

def process_cv_file(input_path: str, output_path: Optional[str] = None,
                   structured: bool = True) -> dict:
    """
    Enhanced CV processing with structured output option.

    Returns:
        Dictionary containing:
        - success: Boolean
        - method: Extraction method used
        - data: Structured data if structured=True, else text
        - error: Any error message
    """
    result = {
        'input_path': input_path,
        'success': False,
        'method': None,
        'data': None,
        'error': None
    }

    try:
        if structured:
            result['data'] = extract_text(input_path, structured=True)
        else:
            result['data'] = extract_text(input_path, structured=False)

        if result['data']:
            result['success'] = True
            result['method'] = 'structured' if structured else 'direct'

            if output_path:
                try:
                    with open(output_path, "w", encoding="utf-8") as f:
                        if isinstance(result['data'], dict):
                            import json
                            json.dump(result['data'], f, indent=2)
                        else:
                            f.write(result['data'])
                    result['output_path'] = output_path
                except Exception as e:
                    result['error'] = f"Failed to save output: {str(e)}"
                    logging.error(result['error'])
    except Exception as e:
        result['error'] = str(e)
        logging.error(f"Processing failed for {input_path}: {str(e)}")

    return result

if __name__ == "__main__":
    # Example usage
    input_file = "VuHoangLan_CV.pdf"
    output_file = "structured_cv_output.txt"

    # Get structured output (default)
    result = process_cv_file(input_file, output_file)

    if result['success']:
        print(f"Extraction successful using {result['method']} method.")
        if 'output_path' in result:
            print(f"Structured data saved to {result['output_path']}")

        # Sample of accessing structured data
        if isinstance(result['data'], dict):
            print("\nExtracted Metadata:")
            for key, value in result['data'].get('metadata', {}).items():
                print(f"{key.title()}: {value}")
    else:
        print(f"Failed to extract: {result.get('error', 'Unknown error')}")